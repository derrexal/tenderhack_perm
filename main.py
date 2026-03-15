from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
import threading
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Mm, Pt
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field, field_validator

DATA_DIR = Path(__file__).parent / "raw_data"
CONTRACTS_FILE = DATA_DIR / "TenderHack_Контракты_20260313.xlsx"
STE_FILE = DATA_DIR / "TenderHack_СТЕ_20260313.xlsx"
TEMPLATE_PATH = Path(__file__).parent / "report_template.docx"

NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
SAMPLE_LIMIT = 5
DEFAULT_CURRENCY = "RUB"
TABLE_ROW_HEIGHT_CM = 0.7
PLACEHOLDER_RE = re.compile(r"\{[^}]+\}")


class Ste(BaseModel):
    id: Optional[int] = None
    name: Optional[str] = None
    category: Optional[str] = None
    manufacturer: Optional[str] = None
    characteristics: Optional[str] = None


class Contract(BaseModel):
    id: Optional[int] = None
    procurementName: Optional[str] = None
    procurementMethod: Optional[str] = None
    initialContractValue: Optional[str] = None
    contractValueAfterSigning: Optional[str] = None
    reductionPercent: Optional[str] = None
    vatRate: Optional[str] = None
    contractSigningDate: Optional[str] = None
    buyerInn: Optional[str] = None
    buyerRegion: Optional[str] = None
    supplierInn: Optional[str] = None
    supplierRegion: Optional[str] = None


class ContractItemWithSte(BaseModel):
    id: Optional[int] = None
    steId: Optional[int] = None
    steItemName: Optional[str] = None
    quantity: Optional[str] = None
    unit: Optional[str] = None
    unitPrice: Optional[str] = None
    steName: Optional[str] = None
    steCategory: Optional[str] = None
    steManufacturer: Optional[str] = None
    steCharacteristics: Optional[str] = None


class JustificationRequest(BaseModel):
    contract: Contract
    items: List[ContractItemWithSte]
    currency: str = Field(default=DEFAULT_CURRENCY, description="Код валюты, по умолчанию RUB")
    maxSamples: int = Field(default=3, description="Сколько примеров контрактов показывать в документе")


class ContractReport(BaseModel):
    contract: Contract
    items: List[ContractItemWithSte]


class BatchJustificationRequest(BaseModel):
    contracts: List[ContractReport]
    currency: str = Field(default=DEFAULT_CURRENCY, description="Код валюты, по умолчанию RUB")
    maxSamples: int = Field(default=3, description="Сколько примеров контрактов показывать в документе")
    reportTitle: Optional[str] = Field(
        default=None, description="Заголовок отчета. Если не задан, используется стандартный."
    )


class StePriceJustificationRow(BaseModel):
    contractId: Optional[str] = None
    contractSigningDate: Optional[str] = None
    buyerInn: Optional[str] = None
    buyerRegion: Optional[str] = None
    count: Optional[str] = None
    unit: Optional[str] = None
    steId: Optional[int] = None
    steName: Optional[str] = None
    unitPrice: Optional[str] = None
    nds: Optional[str] = None

    @field_validator("count", mode="before")
    @classmethod
    def _cast_qty_to_str(cls, value):
        if value is None:
            return None
        return str(value)

    @field_validator("nds", mode="before")
    @classmethod
    def _cast_nds_to_str(cls, value):
        if value is None:
            return None
        return str(value)

    @field_validator("unitPrice", mode="before")
    @classmethod
    def _cast_unit_price_to_str(cls, value):
        if value is None:
            return None
        return str(value)


class StePositionPayload(BaseModel):
    positionName: str = ""
    positionPrice: Optional[Decimal] = None
    items: List[StePriceJustificationRow]


class StePriceTemplateRequest(BaseModel):
    contractName: str = ""
    summaryPrice: Decimal
    positions: List[StePositionPayload]
    currency: str = Field(default=DEFAULT_CURRENCY, description="Код валюты, по умолчанию RUB")
    docType: Optional[str] = Field(
        default=None, description="Формат документа: docx, doc, pdf. По умолчанию docx."
    )


@dataclass
class PriceSample:
    contract_id: str
    signed_at: datetime
    unit_price: Decimal


@dataclass
class PriceStats:
    count: int = 0
    sum_price: Decimal = Decimal("0")
    min_price: Optional[Decimal] = None
    max_price: Optional[Decimal] = None
    min_date: Optional[datetime] = None
    max_date: Optional[datetime] = None
    samples: List[PriceSample] = None

    def __post_init__(self) -> None:
        if self.samples is None:
            self.samples = []

    def update(self, price: Decimal, signed_at: Optional[datetime], contract_id: Optional[str]) -> None:
        self.count += 1
        self.sum_price += price
        if self.min_price is None or price < self.min_price:
            self.min_price = price
        if self.max_price is None or price > self.max_price:
            self.max_price = price
        if signed_at:
            if self.min_date is None or signed_at < self.min_date:
                self.min_date = signed_at
            if self.max_date is None or signed_at > self.max_date:
                self.max_date = signed_at
        if signed_at and contract_id:
            self.samples.append(PriceSample(contract_id=contract_id, signed_at=signed_at, unit_price=price))
            self.samples.sort(key=lambda s: s.signed_at, reverse=True)
            if len(self.samples) > SAMPLE_LIMIT:
                self.samples = self.samples[:SAMPLE_LIMIT]

    @property
    def avg_price(self) -> Optional[Decimal]:
        if not self.count:
            return None
        return self.sum_price / Decimal(self.count)


@dataclass
class ContractReportData:
    contract: Contract
    items: List[ContractItemWithSte]
    analysis: List[Dict[str, object]]
    total_value: Decimal


class RawDataRepository:
    def __init__(self, contracts_path: Path, ste_path: Path) -> None:
        self.contracts_path = contracts_path
        self.ste_path = ste_path
        self._loaded = False
        self._lock = threading.Lock()
        self.ste_by_id: Dict[str, Ste] = {}
        self.stats_by_ste_id: Dict[str, PriceStats] = {}
        self.stats_by_ste_id_method: Dict[Tuple[str, str], PriceStats] = {}
        self.stats_by_name: Dict[str, PriceStats] = {}
        self.stats_by_category: Dict[str, PriceStats] = {}

    def ensure_loaded(self) -> None:
        if self._loaded:
            return
        with self._lock:
            if self._loaded:
                return
            if not self.contracts_path.exists() or not self.ste_path.exists():
                raise FileNotFoundError("raw_data файлы не найдены")
            self.ste_by_id = load_ste_catalog(self.ste_path)
            (
                self.stats_by_ste_id,
                self.stats_by_ste_id_method,
                self.stats_by_name,
                self.stats_by_category,
            ) = load_contract_stats(self.contracts_path, self.ste_by_id)
            self._loaded = True


DATA_REPO = RawDataRepository(CONTRACTS_FILE, STE_FILE)


def normalize_text(value: Optional[str]) -> str:
    if not value:
        return ""
    text = value.lower().strip()
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_name(value: Optional[str]) -> str:
    if not value:
        return ""
    text = value.lower()
    text = re.sub(r"[^0-9a-zа-яё\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def parse_decimal(value: Optional[str]) -> Optional[Decimal]:
    if value is None:
        return None
    if isinstance(value, Decimal):
        return value
    if isinstance(value, (int, float)):
        return Decimal(str(value))
    text = str(value).strip()
    if not text:
        return None
    text = text.replace(" ", "").replace(",", ".")
    text = re.sub(r"[^0-9.\-]", "", text)
    if not text or text in {"-", "."}:
        return None
    try:
        return Decimal(text)
    except Exception:
        return None


def parse_datetime(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    text = str(value).strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S.%f", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    try:
        return datetime.fromisoformat(text)
    except ValueError:
        return None


def format_decimal(value: Optional[Decimal], places: int = 2) -> str:
    if value is None:
        return "—"
    quant = Decimal("1") if places == 0 else Decimal("0." + "0" * (places - 1) + "1")
    amount = value.quantize(quant, rounding=ROUND_HALF_UP)
    formatted = f"{amount:,.{places}f}".replace(",", " ")
    return formatted.replace(".", ",")


def format_optional_decimal(value: Optional[str], places: int = 2) -> str:
    if value is None:
        return ""
    parsed = parse_decimal(value)
    if parsed is None:
        return str(value)
    return format_decimal(parsed, places)


def format_optional_date(value: Optional[str]) -> str:
    if not value:
        return ""
    parsed = parse_datetime(value)
    if parsed is None:
        return str(value)
    return parsed.strftime("%d.%m.%Y")


def _ru_plural(value: int, one: str, few: str, many: str) -> str:
    value = abs(value) % 100
    if 11 <= value <= 19:
        return many
    last = value % 10
    if last == 1:
        return one
    if 2 <= last <= 4:
        return few
    return many


def _ru_triad_to_words(value: int, gender: str) -> List[str]:
    ones_m = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = [
        "десять",
        "одиннадцать",
        "двенадцать",
        "тринадцать",
        "четырнадцать",
        "пятнадцать",
        "шестнадцать",
        "семнадцать",
        "восемнадцать",
        "девятнадцать",
    ]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    words: List[str] = []
    if value == 0:
        return words
    h = value // 100
    if h:
        words.append(hundreds[h])
    t = value % 100
    if 10 <= t <= 19:
        words.append(teens[t - 10])
        return words
    ten = t // 10
    if ten:
        words.append(tens[ten])
    unit = t % 10
    if unit:
        words.append((ones_f if gender == "f" else ones_m)[unit])
    return words


def _ru_number_to_words(value: int) -> str:
    if value == 0:
        return "ноль"
    scales = [
        (1_000_000_000, "миллиард", "миллиарда", "миллиардов", "m"),
        (1_000_000, "миллион", "миллиона", "миллионов", "m"),
        (1000, "тысяча", "тысячи", "тысяч", "f"),
    ]
    words: List[str] = []
    remainder = value
    for scale_value, one, few, many, gender in scales:
        part = remainder // scale_value
        if part:
            words.extend(_ru_triad_to_words(part, gender))
            words.append(_ru_plural(part, one, few, many))
            remainder %= scale_value
    words.extend(_ru_triad_to_words(remainder, "m"))
    return " ".join(words).strip()


def amount_to_rubles_words(amount: Decimal) -> str:
    quantized = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rubles = int(quantized)
    kopeks = int((quantized - Decimal(rubles)) * 100)
    rubles_words = _ru_number_to_words(rubles)
    rubles_word = _ru_plural(rubles, "рубль", "рубля", "рублей")
    kopeks_word = _ru_plural(kopeks, "копейка", "копейки", "копеек")
    return f"{rubles_words} {rubles_word} {kopeks:02d} {kopeks_word}"


def _replace_placeholders_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    if "{" not in full_text:
        return
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    run_spans = []
    pos = 0
    for run in paragraph.runs:
        start = pos
        end = pos + len(run.text)
        run_spans.append((run, start, end))
        pos = end

    for match in reversed(matches):
        key = match.group(0)[1:-1].strip()
        if key not in mapping:
            continue
        replacement = mapping.get(key, "")

        first = None
        last = None
        for idx, (run, start, end) in enumerate(run_spans):
            if end <= match.start():
                continue
            if start >= match.end():
                break
            if first is None:
                first = idx
            last = idx

        if first is None or last is None:
            continue

        first_run, f_start, f_end = run_spans[first]
        last_run, l_start, l_end = run_spans[last]

        prefix = first_run.text[: max(0, match.start() - f_start)]
        suffix = last_run.text[max(0, match.end() - l_start) :]

        if first == last:
            first_run.text = prefix + replacement + suffix
        else:
            first_run.text = prefix + replacement
            last_run.text = suffix
            for idx in range(first + 1, last):
                run_spans[idx][0].text = ""


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    if not paragraph.text:
        return
    _replace_placeholders_in_paragraph(paragraph, replacements)


def _replace_placeholders(doc: Document, replacements: Dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _replace_placeholders_sequence(doc: Document, replacements: Dict[str, List[str]]) -> None:
    counters = {key: 0 for key in replacements}

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, values in replacements.items():
                while key in run.text:
                    idx = counters[key]
                    value = values[idx] if idx < len(values) else ""
                    run.text = run.text.replace(key, value, 1)
                    counters[key] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, values in replacements.items():
                            while key in run.text:
                                idx = counters[key]
                                value = values[idx] if idx < len(values) else ""
                                run.text = run.text.replace(key, value, 1)
                                counters[key] += 1


def _collect_placeholders_from_runs(runs) -> List[str]:
    text = "".join(run.text or "" for run in runs)
    keys = []
    for match in PLACEHOLDER_RE.findall(text):
        keys.append(match[1:-1].strip())
    return keys


def _collect_placeholders_from_paragraph(paragraph) -> List[str]:
    return _collect_placeholders_from_runs(paragraph.runs)


def _collect_placeholders_from_table(table) -> List[str]:
    keys: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                keys.extend(_collect_placeholders_from_paragraph(paragraph))
    return keys


def _collect_placeholders_from_doc(doc: Document) -> List[str]:
    keys: List[str] = []
    for paragraph in doc.paragraphs:
        keys.extend(_collect_placeholders_from_paragraph(paragraph))
    for table in doc.tables:
        keys.extend(_collect_placeholders_from_table(table))
    return keys


def _is_items_table(table) -> bool:
    if not table.rows:
        return False
    header = [cell.text.strip().lower() for cell in table.rows[0].cells]
    normalized = [h.replace("ё", "е") for h in header]
    return len(normalized) >= 3 and "№" in normalized[0] and "идентификатор сте" in normalized[1]


def _normalize_header(text: str) -> str:
    cleaned = text.replace("ё", "е").lower()
    cleaned = re.sub(r"\\s+", " ", cleaned).strip()
    return cleaned


def _fill_items_table(
    table,
    items: List[StePriceJustificationRow],
    position_price: Optional[Decimal],
) -> None:
    if not table.rows:
        raise ValueError("Таблица в шаблоне пуста.")

    item_row = None
    summary_row = None
    for row in table.rows:
        keys = _collect_placeholders_from_runs(
            [run for cell in row.cells for p in cell.paragraphs for run in p.runs]
        )
        if keys and "positionPrice" not in keys:
            item_row = row
        if "positionPrice" in keys:
            summary_row = row

    if item_row is None:
        raise ValueError("В шаблоне не найдена строка с плейсхолдерами для позиций.")
    if summary_row is None:
        raise ValueError("В шаблоне не найдена строка с плейсхолдером positionPrice.")

    keep_trs = {table.rows[0]._tr, item_row._tr, summary_row._tr}
    for row in list(table.rows)[::-1]:
        if row._tr not in keep_trs:
            table._tbl.remove(row._tr)
    table.__dict__.pop("rows", None)

    item_row_template = deepcopy(item_row._tr)

    item_placeholders = set(
        _collect_placeholders_from_runs(
            [run for cell in item_row.cells for p in cell.paragraphs for run in p.runs]
        )
    )

    def _set_row_index(row, index_value: int) -> None:
        if not row.cells:
            return
        cell = row.cells[0]
        if cell.paragraphs:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            cell.paragraphs[0].text = str(index_value)
        else:
            cell.text = str(index_value)

    optional_placeholders = {"steName", "nds"}

    def item_mapping(item: StePriceJustificationRow) -> Dict[str, str]:
        values = {
            "steId": str(item.steId) if item.steId is not None else "",
            "contractId": item.contractId or "",
            "contractSigningDate": format_optional_date(item.contractSigningDate),
            "buyerInn": item.buyerInn or "",
            "count": item.count or "",
            "unit": item.unit or "",
            "buyerRegion": item.buyerRegion or "",
            "unitPrice": format_optional_decimal(item.unitPrice),
        }
        # steName и nds не заполняются даже если пришли в payload.
        if "steName" in item_placeholders:
            values["steName"] = ""
        if "nds" in item_placeholders:
            values["nds"] = ""
        return {key: values.get(key, "") for key in item_placeholders}

    def validate_item_placeholders(mapping: Dict[str, str]) -> None:
        missing = [
            k
            for k in item_placeholders
            if mapping.get(k, "") == "" and k not in optional_placeholders
        ]
        extra = [k for k, v in mapping.items() if v and k not in item_placeholders]
        if missing or extra:
            parts = []
            if missing:
                parts.append(f"не переданы: {', '.join(sorted(missing))}")
            if extra:
                parts.append(f"лишние: {', '.join(sorted(extra))}")
            raise ValueError("Некорректные поля для строки таблицы: " + "; ".join(parts))

    mapping_first = item_mapping(items[0])
    validate_item_placeholders(mapping_first)
    for cell in item_row.cells:
        for paragraph in cell.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, mapping_first)
    _set_row_index(item_row, 1)

    summary_keys = set(
        _collect_placeholders_from_runs(
            [run for cell in summary_row.cells for p in cell.paragraphs for run in p.runs]
        )
    )
    summary_mapping = {
        "positionPrice": format_decimal(position_price) if position_price is not None else ""
    }
    if "positionPrice" in summary_keys and summary_mapping["positionPrice"] == "":
        raise ValueError("positionPrice обязателен, так как используется в шаблоне.")
    if "positionPrice" not in summary_keys and summary_mapping["positionPrice"] != "":
        raise ValueError("positionPrice передан, но отсутствует в шаблоне.")
    for cell in summary_row.cells:
        for paragraph in cell.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, summary_mapping)

    summary_tr = summary_row._tr
    for item_idx, item in enumerate(items[1:], start=2):
        new_tr = deepcopy(item_row_template)
        summary_tr.addprevious(new_tr)
        table.__dict__.pop("rows", None)
        new_row = _Row(new_tr, table)
        mapping = item_mapping(item)
        validate_item_placeholders(mapping)
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, mapping)
        _set_row_index(new_row, item_idx)


def _clear_items_table(table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)


def _set_position_heading_text(paragraph: Paragraph, position_name: str) -> None:
    _replace_in_paragraph(paragraph, {"предмет закупки": position_name or ""})


def currency_label(code: str) -> str:
    code_up = (code or "").upper()
    if code_up in {"RUB", "RUR", "РУБ", "RUBLE"}:
        return "руб."
    return code


def rtf_escape(text: str) -> str:
    escaped = []
    for ch in text:
        if ch in {"\\", "{", "}"}:
            escaped.append("\\" + ch)
        elif ch == "\n":
            escaped.append("\\par ")
        elif ord(ch) < 128:
            escaped.append(ch)
        else:
            escaped.append(f"\\u{ord(ch)}?")
    return "".join(escaped)


def _ensure_rfonts(style, font_name: str) -> None:
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _ensure_paragraph_style(
    styles,
    name: str,
    font_name: str,
    font_size: int,
    bold: bool = False,
    alignment: Optional[int] = None,
    line_spacing: Optional[float] = None,
    first_line_indent_cm: Optional[float] = None,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    based_on: Optional[str] = None,
) -> None:
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if based_on:
        style.base_style = styles[based_on]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    _ensure_rfonts(style, font_name)

    pf = style.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


def _ensure_table_style(styles, name: str, base: str = "Table Grid") -> None:
    if name in styles:
        return
    style = styles.add_style(name, WD_STYLE_TYPE.TABLE)
    style.base_style = styles[base]
    style.hidden = False
    style.quick_style = True
    style.priority = 99

    tbl_style_pr = OxmlElement("w:tblStylePr")
    tbl_style_pr.set(qn("w:type"), "firstRow")

    r_pr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    r_pr.append(b)
    tbl_style_pr.append(r_pr)

    tc_pr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E6E6E6")
    tc_pr.append(shd)
    tbl_style_pr.append(tc_pr)

    style._element.append(tbl_style_pr)


def apply_uniform_font_and_table_formatting(doc: Document, font_name: str = "Times New Roman") -> None:
    return


def _set_repeat_table_header(row) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _find_position_blocks(doc: Document) -> List[Tuple[Paragraph, Table]]:
    blocks: List[Tuple[Paragraph, Table]] = []
    current_paragraph: Optional[Paragraph] = None
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            if "Расчет начальной (максимальной) цены" in (block.text or ""):
                current_paragraph = block
        elif isinstance(block, Table) and current_paragraph is not None:
            if _is_items_table(block):
                blocks.append((current_paragraph, block))
                current_paragraph = None
    return blocks


def _insert_after(element, new_element) -> None:
    element.addnext(new_element)


def _insert_before(element, new_element) -> None:
    element.addprevious(new_element)


def _get_message_paragraph_template(doc: Document) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name in {"Основной текст1", "Основной текст"}:
            return paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            return paragraph
    raise ValueError("В шаблоне не найден подходящий абзац для вставки сообщения.")


def _build_no_items_message(position: StePositionPayload) -> str:
    if not position.positionName:
        raise ValueError("positionName обязателен, так как items пуст.")
    if position.positionPrice is None:
        raise ValueError("positionPrice обязателен, так как items пуст.")
    price_text = format_decimal(position.positionPrice)
    return (
        f"Данных для рассчета НПМЦ товара {position.positionName} недостаточно, "
        f"выставлено ручное значение {price_text}."
    )



class RtfBuilder:
    def __init__(self, landscape: bool = False, margin_twips: int = 1134) -> None:
        if landscape:
            paperw, paperh = 16840, 11907
        else:
            paperw, paperh = 11907, 16840
        self.parts: List[str] = [
            "{\\rtf1\\ansi\\deff0"
            "{\\fonttbl{\\f0 Calibri;}}"
            "{\\colortbl ;\\red230\\green230\\blue230;}"
            f"\\paperw{paperw}\\paperh{paperh}"
            f"\\margl{margin_twips}\\margr{margin_twips}"
            f"\\margt{margin_twips}\\margb{margin_twips}"
            "\\viewkind4\\uc1\\pard\\fs22 "
        ]

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        size: Optional[int] = None,
        align: str = "left",
        space_after: int = 120,
    ) -> None:
        prefix = ""
        suffix = ""
        align_code = {"left": "\\ql", "center": "\\qc", "right": "\\qr", "justify": "\\qj"}.get(
            align, "\\ql"
        )
        if bold:
            prefix += "\\b "
            suffix += "\\b0 "
        if size is not None:
            prefix += f"\\fs{size} "
            suffix += "\\fs22 "
        spacing = f"\\sa{space_after} "
        self.parts.append(
            f"{align_code}{spacing}" + prefix + rtf_escape(text) + suffix + "\\par "
        )

    def finish(self) -> bytes:
        self.parts.append("}")
        return "".join(self.parts).encode("ascii", errors="ignore")

    def add_page_break(self) -> None:
        self.parts.append("\\page ")

    def add_table(
        self,
        rows: List[List[str]],
        col_widths: List[int],
        header_rows: int = 1,
        header_shading: bool = True,
        borders: bool = True,
        font_size: int = 20,
    ) -> None:
        if not rows or not col_widths:
            return
        col_positions = []
        acc = 0
        for width in col_widths:
            acc += width
            col_positions.append(acc)

        total_cols = len(col_widths)
        border_defs = ""
        if borders:
            border_defs = "\\clbrdrt\\brdrs\\clbrdrl\\brdrs\\clbrdrb\\brdrs\\clbrdrr\\brdrs"

        for row_idx, row in enumerate(rows):
            padded_row = row + [""] * max(0, total_cols - len(row))
            row_rtf = ["\\trowd\\trgaph108\\trleft0"]
            if row_idx < header_rows:
                row_rtf.append("\\trhdr")
            for pos in col_positions:
                cell_defs = border_defs
                if header_shading and row_idx < header_rows:
                    cell_defs += "\\clcbpat1"
                row_rtf.append(f"{cell_defs}\\cellx{pos}")
            row_rtf.append(" ")
            for cell in padded_row[:total_cols]:
                cell_text = rtf_escape(cell)
                if row_idx < header_rows:
                    row_rtf.append(
                        f"\\intbl\\qc\\b\\fs{font_size} {cell_text}\\b0\\fs22\\cell "
                    )
                else:
                    row_rtf.append(
                        f"\\intbl\\ql\\fs{font_size} {cell_text}\\fs22\\cell "
                    )
            row_rtf.append("\\row ")
            self.parts.append("".join(row_rtf))


def col_to_idx(col: str) -> int:
    idx = 0
    for ch in col:
        idx = idx * 26 + (ord(ch) - 64)
    return idx - 1


def load_shared_strings(z: zipfile.ZipFile) -> List[str]:
    if "xl/sharedStrings.xml" not in z.namelist():
        return []
    strings: List[str] = []
    with z.open("xl/sharedStrings.xml") as f:
        for _, elem in ET.iterparse(f, events=("end",)):
            if elem.tag == f"{NS}si":
                text = "".join(
                    t.text or "" for t in elem.findall(f".//{NS}t")
                )
                strings.append(text)
                elem.clear()
    return strings


def iter_xlsx_rows(path: Path) -> Iterable[List[str]]:
    with zipfile.ZipFile(path) as z:
        shared = load_shared_strings(z)
        with z.open("xl/worksheets/sheet1.xml") as f:
            context = ET.iterparse(f, events=("end",))
            for _, elem in context:
                if elem.tag != f"{NS}row":
                    continue
                row_cells: Dict[int, str] = {}
                for cell in elem.findall(f"{NS}c"):
                    ref = cell.get("r")
                    if not ref:
                        continue
                    col = "".join(ch for ch in ref if ch.isalpha())
                    col_idx = col_to_idx(col)
                    cell_type = cell.get("t")
                    value = ""
                    if cell_type == "inlineStr":
                        is_elem = cell.find(f"{NS}is")
                        if is_elem is not None:
                            value = "".join(
                                t.text or "" for t in is_elem.findall(f".//{NS}t")
                            )
                    else:
                        v = cell.find(f"{NS}v")
                        if v is not None and v.text is not None:
                            value = v.text
                            if cell_type == "s":
                                try:
                                    value = shared[int(value)]
                                except Exception:
                                    pass
                    row_cells[col_idx] = value
                if row_cells:
                    max_idx = max(row_cells)
                    row = [""] * (max_idx + 1)
                    for idx, val in row_cells.items():
                        row[idx] = val
                else:
                    row = []
                yield row
                elem.clear()


def safe_get(row: List[str], idx: Optional[int]) -> str:
    if idx is None:
        return ""
    if idx >= len(row):
        return ""
    return row[idx]


def load_ste_catalog(path: Path) -> Dict[str, Ste]:
    rows = iter_xlsx_rows(path)
    header = next(rows, [])
    index = {name: idx for idx, name in enumerate(header)}
    ste_map: Dict[str, Ste] = {}
    for row in rows:
        ste_id = safe_get(row, index.get("Идентификатор СТЕ"))
        if not ste_id:
            continue
        ste = Ste(
            id=int(ste_id) if str(ste_id).isdigit() else None,
            name=safe_get(row, index.get("Наименование СТЕ")) or None,
            category=safe_get(row, index.get("Категория")) or None,
            manufacturer=safe_get(row, index.get("Производитель")) or None,
            characteristics=safe_get(row, index.get("характеристики СТЕ")) or None,
        )
        ste_map[str(ste_id)] = ste
    return ste_map


def load_contract_stats(
    path: Path,
    ste_map: Dict[str, Ste],
) -> Tuple[Dict[str, PriceStats], Dict[Tuple[str, str], PriceStats], Dict[str, PriceStats], Dict[str, PriceStats]]:
    rows = iter_xlsx_rows(path)
    header = next(rows, [])
    index = {name: idx for idx, name in enumerate(header)}
    stats_by_ste: Dict[str, PriceStats] = {}
    stats_by_ste_method: Dict[Tuple[str, str], PriceStats] = {}
    stats_by_name: Dict[str, PriceStats] = {}
    stats_by_category: Dict[str, PriceStats] = {}

    for row in rows:
        ste_id = safe_get(row, index.get("Идентификатор СТЕ по контракту"))
        ste_name = safe_get(row, index.get("Наименование позиции СТЕ"))
        unit_price_raw = safe_get(row, index.get("Цена за единицу"))
        if not unit_price_raw:
            continue
        unit_price = parse_decimal(unit_price_raw)
        if unit_price is None:
            continue
        contract_id = safe_get(row, index.get("Идентификатор контракта"))
        signed_at = parse_datetime(safe_get(row, index.get("Дата заключения контракта")))
        procurement_method = safe_get(row, index.get("Способ закупки"))

        if ste_id:
            ste_key = str(ste_id)
            stats_by_ste.setdefault(ste_key, PriceStats()).update(
                unit_price, signed_at, contract_id
            )
            if procurement_method:
                method_key = (ste_key, procurement_method)
                stats_by_ste_method.setdefault(method_key, PriceStats()).update(
                    unit_price, signed_at, contract_id
                )
            ste_info = ste_map.get(ste_key)
            if ste_info and ste_info.category:
                cat_key = normalize_text(ste_info.category)
                if cat_key:
                    stats_by_category.setdefault(cat_key, PriceStats()).update(
                        unit_price, signed_at, contract_id
                    )

        if ste_name:
            name_key = normalize_name(ste_name)
            if name_key:
                stats_by_name.setdefault(name_key, PriceStats()).update(
                    unit_price, signed_at, contract_id
                )

    return stats_by_ste, stats_by_ste_method, stats_by_name, stats_by_category


def enrich_item(item: ContractItemWithSte, ste_map: Dict[str, Ste]) -> ContractItemWithSte:
    if item.steId is None:
        return item
    ste = ste_map.get(str(item.steId))
    if not ste:
        return item
    if not item.steName:
        item.steName = ste.name
    if not item.steCategory:
        item.steCategory = ste.category
    if not item.steManufacturer:
        item.steManufacturer = ste.manufacturer
    if not item.steCharacteristics:
        item.steCharacteristics = ste.characteristics
    return item


def choose_stats(
    item: ContractItemWithSte,
    repo: RawDataRepository,
    procurement_method: Optional[str],
) -> Tuple[Optional[PriceStats], str]:
    if item.steId is not None:
        ste_key = str(item.steId)
        if procurement_method:
            stats = repo.stats_by_ste_id_method.get((ste_key, procurement_method))
            if stats and stats.count:
                return stats, "СТЕ + способ закупки"
        stats = repo.stats_by_ste_id.get(ste_key)
        if stats and stats.count:
            return stats, "СТЕ"
    name_key = normalize_name(item.steItemName or item.steName)
    if name_key:
        stats = repo.stats_by_name.get(name_key)
        if stats and stats.count:
            return stats, "наименование позиции"
    cat_key = normalize_text(item.steCategory)
    if cat_key:
        stats = repo.stats_by_category.get(cat_key)
        if stats and stats.count:
            return stats, "категория"
    return None, "нет данных"


def analyze_contract(
    contract: Contract,
    items: List[ContractItemWithSte],
    repo: RawDataRepository,
    max_samples: int,
) -> ContractReportData:
    enriched_items: List[ContractItemWithSte] = [
        enrich_item(item, repo.ste_by_id) for item in items
    ]

    analysis: List[Dict[str, object]] = []
    total_value = Decimal("0")
    sample_limit = max(0, min(max_samples, SAMPLE_LIMIT))

    for item in enriched_items:
        stats, basis = choose_stats(item, repo, contract.procurementMethod)
        quantity = parse_decimal(item.quantity) or Decimal("0")
        unit = item.unit or ""
        fallback_price = parse_decimal(item.unitPrice) or Decimal("0")
        if stats and stats.avg_price is not None:
            unit_price = stats.avg_price
        else:
            unit_price = fallback_price
        cost = quantity * unit_price
        total_value += cost

        analysis.append(
            {
                "name": item.steItemName or item.steName or "Без наименования",
                "quantity": quantity,
                "unit": unit,
                "unit_price": unit_price,
                "cost": cost,
                "stats": stats,
                "basis": basis,
                "fallback": format_decimal(fallback_price),
                "max_samples": sample_limit,
            }
        )

    return ContractReportData(
        contract=contract,
        items=enriched_items,
        analysis=analysis,
        total_value=total_value,
    )


def append_procurement_section(builder: RtfBuilder, contract: Contract, title: str) -> None:
    builder.add_paragraph(title, bold=True)
    if contract.procurementName:
        builder.add_paragraph(f"Наименование закупки: {contract.procurementName}")
    if contract.procurementMethod:
        builder.add_paragraph(f"Способ закупки: {contract.procurementMethod}")
    if contract.id is not None:
        builder.add_paragraph(f"Идентификатор контракта: {contract.id}")
    if contract.buyerInn or contract.buyerRegion:
        buyer = " ".join(part for part in [contract.buyerInn, contract.buyerRegion] if part)
        builder.add_paragraph(f"Заказчик: {buyer}")
    if contract.supplierInn or contract.supplierRegion:
        supplier = " ".join(
            part for part in [contract.supplierInn, contract.supplierRegion] if part
        )
        builder.add_paragraph(f"Поставщик (при наличии): {supplier}")
    if contract.contractSigningDate:
        builder.add_paragraph(f"Дата заключения контракта: {contract.contractSigningDate}")
    if contract.vatRate:
        builder.add_paragraph(f"Ставка НДС: {contract.vatRate}")


def append_items_section(builder: RtfBuilder, items: List[ContractItemWithSte], title: str) -> None:
    builder.add_paragraph(title, bold=True)
    for idx, item in enumerate(items, start=1):
        name = item.steItemName or item.steName or "Без наименования"
        qty = parse_decimal(item.quantity) or Decimal("0")
        unit = item.unit or ""
        ste_id = item.steId if item.steId is not None else "—"
        builder.add_paragraph(f"{idx}. {name}")
        builder.add_paragraph(f"Количество: {format_decimal(qty, 3)} {unit}. СТЕ: {ste_id}.")
        if item.steCategory or item.steManufacturer:
            cat = item.steCategory or "—"
            mfr = item.steManufacturer or "—"
            builder.add_paragraph(f"Категория: {cat}. Производитель: {mfr}.")
        if item.steCharacteristics:
            builder.add_paragraph(f"Характеристики: {item.steCharacteristics}")


def append_market_section(
    builder: RtfBuilder,
    analysis: List[Dict[str, object]],
    currency: str,
    title: str,
) -> None:
    builder.add_paragraph(title, bold=True)
    for idx, item_info in enumerate(analysis, start=1):
        builder.add_paragraph(f"{idx}. {item_info['name']}")
        basis = item_info["basis"]
        stats: Optional[PriceStats] = item_info["stats"]
        if stats and stats.count:
            period = ""
            if stats.min_date and stats.max_date:
                period = (
                    f" (период {stats.min_date.strftime('%d.%m.%Y')}"
                    f" — {stats.max_date.strftime('%d.%m.%Y')})"
                )
            builder.add_paragraph(
                f"Найдено сопоставимых контрактов: {stats.count}{period}. Основание: {basis}."
            )
            builder.add_paragraph(
                "Минимальная цена: "
                f"{format_decimal(stats.min_price)} {currency_label(currency)}; "
                "Средняя цена: "
                f"{format_decimal(stats.avg_price)} {currency_label(currency)}; "
                "Максимальная цена: "
                f"{format_decimal(stats.max_price)} {currency_label(currency)}."
            )
            samples = stats.samples[: item_info["max_samples"]]
            if samples:
                sample_lines = []
                for sample in samples:
                    sample_lines.append(
                        f"№{sample.contract_id} от {sample.signed_at.strftime('%d.%m.%Y')}: "
                        f"{format_decimal(sample.unit_price)} {currency_label(currency)}"
                    )
                builder.add_paragraph("Примеры контрактов: " + "; ".join(sample_lines))
        else:
            fallback = item_info["fallback"]
            builder.add_paragraph(
                "Сопоставимые контракты не найдены. "
                f"В качестве ориентира использована цена: {fallback} {currency_label(currency)}."
            )


def append_calc_section(
    builder: RtfBuilder,
    analysis: List[Dict[str, object]],
    total_value: Decimal,
    currency: str,
    title: str,
    total_label: str,
    total_bold: bool = True,
) -> None:
    builder.add_paragraph(title, bold=True)
    for idx, item_info in enumerate(analysis, start=1):
        qty = item_info["quantity"]
        unit = item_info["unit"]
        unit_price = item_info["unit_price"]
        cost = item_info["cost"]
        builder.add_paragraph(
            f"{idx}. {item_info['name']}: {format_decimal(qty, 3)} {unit} × "
            f"{format_decimal(unit_price)} = {format_decimal(cost)} {currency_label(currency)}"
        )
    builder.add_paragraph(
        f"{total_label}{format_decimal(total_value)} {currency_label(currency)}",
        bold=total_bold,
    )


def append_conclusion(builder: RtfBuilder, title: str) -> None:
    builder.add_paragraph(title, bold=True)
    builder.add_paragraph(
        "На основании анализа данных о ранее заключенных контрактах и справочника СТЕ "
        "рекомендуется установить начальную (максимальную) цену контракта на уровне, "
        "указанном в разделе расчета."
    )


def build_document(
    contract: Contract,
    items: List[ContractItemWithSte],
    analysis: List[Dict[str, object]],
    total_value: Decimal,
    currency: str,
) -> bytes:
    builder = RtfBuilder()
    today = date.today().strftime("%d.%m.%Y")

    builder.add_paragraph(
        "Обоснование начальной (максимальной) цены контракта для котировочной сессии",
        bold=True,
        size=32,
    )
    builder.add_paragraph(f"Дата формирования: {today}")

    append_procurement_section(builder, contract, "1. Сведения о закупке")
    append_items_section(builder, items, "2. Объект закупки")
    append_market_section(builder, analysis, currency, "3. Анализ рынка")
    append_calc_section(
        builder,
        analysis,
        total_value,
        currency,
        "4. Расчет НМЦК",
        "Итого НМЦК: ",
    )
    append_conclusion(builder, "5. Вывод")

    return builder.finish()


def build_batch_document(
    reports: List[ContractReportData],
    currency: str,
    report_title: Optional[str],
) -> bytes:
    builder = RtfBuilder()
    today = date.today().strftime("%d.%m.%Y")
    title = report_title or "Сводное обоснование начальной (максимальной) цены контрактов"

    builder.add_paragraph(title, bold=True, size=32)
    builder.add_paragraph(f"Дата формирования: {today}")
    builder.add_paragraph(f"Количество контрактов: {len(reports)}")

    total_all = sum((report.total_value for report in reports), Decimal("0"))
    builder.add_paragraph(
        f"Суммарная НМЦК: {format_decimal(total_all)} {currency_label(currency)}",
        bold=True,
    )

    builder.add_paragraph("1. Сводка по контрактам", bold=True)
    for idx, report in enumerate(reports, start=1):
        contract = report.contract
        name = contract.procurementName or "Без наименования"
        method = contract.procurementMethod or "—"
        items_count = len(report.items)
        builder.add_paragraph(f"{idx}. {name}")
        builder.add_paragraph(
            f"Способ закупки: {method}. Позиции: {items_count}. "
            f"НМЦК: {format_decimal(report.total_value)} {currency_label(currency)}."
        )

    builder.add_paragraph("2. Детализация по контрактам", bold=True)
    for idx, report in enumerate(reports, start=1):
        name = report.contract.procurementName or "Без наименования"
        builder.add_paragraph(f"Контракт {idx}: {name}", bold=True, size=28)
        append_procurement_section(builder, report.contract, "Сведения о закупке")
        append_items_section(builder, report.items, "Объект закупки")
        append_market_section(builder, report.analysis, currency, "Анализ рынка")
        append_calc_section(
            builder,
            report.analysis,
            report.total_value,
            currency,
            "Расчет НМЦК",
            "Итого НМЦК по контракту: ",
        )
        append_conclusion(builder, "Вывод")
        if idx != len(reports):
            builder.add_page_break()

    return builder.finish()


def build_ste_price_document(
    request: StePriceJustificationRequest,
    rows: List[StePriceJustificationRow],
) -> bytes:
    builder = RtfBuilder(landscape=True, margin_twips=720)
    today = date.today().strftime("%d.%m.%Y")
    title = request.reportTitle or "Обоснование начальной цены стандартной товарной единицы"

    ste_id_values = {row.steId for row in rows if row.steId is not None}
    ste_name_values = {row.steItemName for row in rows if row.steItemName}
    ste_id = request.steId or (next(iter(ste_id_values)) if ste_id_values else None)
    ste_name = request.steItemName or (next(iter(ste_name_values)) if ste_name_values else None)

    prices = []
    dates = []
    for row in rows:
        price = parse_decimal(row.unitPrice)
        if price is not None:
            prices.append(price)
        dt = parse_datetime(row.contractSigningDate)
        if dt:
            dates.append(dt)

    count = len(prices)
    avg_price = sum(prices, Decimal("0")) / Decimal(count) if count else None
    min_price = min(prices) if prices else None
    max_price = max(prices) if prices else None
    period = ""
    if dates:
        period = f"{min(dates).strftime('%d.%m.%Y')} — {max(dates).strftime('%d.%m.%Y')}"

    builder.add_paragraph(title, bold=True, size=32, align="center", space_after=240)
    builder.add_paragraph(f"Дата формирования: {today}", align="center", space_after=240)
    info_lines = []
    if ste_id is not None or ste_name:
        ste_line = "СТЕ: "
        if ste_id is not None:
            ste_line += str(ste_id)
        if ste_name:
            ste_line += f" — {ste_name}"
        info_lines.append(ste_line)
    if period:
        info_lines.append(f"Период данных: {period}")
    info_lines.append(f"Количество сопоставимых контрактов: {len(rows)}")
    for line in info_lines:
        builder.add_paragraph(line, align="left", space_after=120)

    if avg_price is not None:
        builder.add_paragraph(
            "Рекомендованная средняя цена: "
            f"{format_decimal(avg_price)} {currency_label(request.currency)}",
            bold=True,
            align="left",
            space_after=240,
        )

    if len(ste_id_values) > 1:
        builder.add_paragraph(
            "Внимание: входные данные содержат несколько идентификаторов СТЕ.",
            bold=True,
            space_after=120,
        )
    if len(ste_name_values) > 1:
        builder.add_paragraph(
            "Внимание: входные данные содержат несколько наименований СТЕ.",
            bold=True,
            space_after=120,
        )

    table_rows: List[List[str]] = [
        [
            "№",
            "ID контракта",
            "Способ закупки",
            "НМЦК",
            "Цена после заключения",
            "% снижения",
            "Дата заключения",
            "ИНН заказчика",
            "ИНН поставщика",
            "СТЕ",
            "Позиция СТЕ",
            "Цена за единицу",
        ]
    ]

    for idx, row in enumerate(rows, start=1):
        table_rows.append(
            [
                str(idx),
                row.contractId or "",
                row.procurementMethod or "",
                format_optional_decimal(row.initialContractValue),
                format_optional_decimal(row.contractValueAfterSigning),
                row.reductionPercent or "",
                format_optional_date(row.contractSigningDate),
                row.buyerInn or "",
                row.supplierInn or "",
                str(row.steId) if row.steId is not None else "",
                row.steItemName or "",
                format_optional_decimal(row.unitPrice),
            ]
        )

    col_widths = [400, 900, 1500, 1000, 1100, 800, 1100, 1100, 1100, 800, 2000, 1000]
    builder.add_table(
        table_rows,
        col_widths,
        header_rows=1,
        header_shading=True,
        borders=True,
        font_size=18,
    )

    if avg_price is not None:
        builder.add_paragraph(
            "Средняя цена: "
            f"{format_decimal(avg_price)} {currency_label(request.currency)}; "
            "Минимальная цена: "
            f"{format_decimal(min_price)} {currency_label(request.currency)}; "
            "Максимальная цена: "
            f"{format_decimal(max_price)} {currency_label(request.currency)}.",
            space_after=240,
        )

    signer_title = request.signerTitle or "Должность"
    signer_name = request.signerName or "ФИО"
    builder.add_paragraph("", space_after=120)
    builder.add_paragraph(
        f"{signer_title}: ____________________ /{signer_name}/",
        align="right",
        space_after=0,
    )

    return builder.finish()


def build_ste_price_docx_from_template(payload: StePriceTemplateRequest) -> bytes:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    today = date.today().strftime("%d.%m.%Y")

    placeholders_in_doc = set(_collect_placeholders_from_doc(doc))
    allowed_placeholders = {
        "наименование закупки",
        "предмет закупки",
        "summaryPrice",
        "сумма русскими словами",
        "today",
        "steId",
        "steName",
        "contractId",
        "contractSigningDate",
        "buyerInn",
        "count",
        "unit",
        "buyerRegion",
        "unitPrice",
        "nds",
        "positionPrice",
    }
    unknown = placeholders_in_doc - allowed_placeholders
    if unknown:
        raise ValueError(
            "Неизвестные плейсхолдеры в шаблоне: " + ", ".join(sorted(unknown))
        )

    replacements = {
        "наименование закупки": payload.contractName or "",
        "summaryPrice": format_decimal(payload.summaryPrice),
        "сумма русскими словами": f"({amount_to_rubles_words(payload.summaryPrice)})",
        "today": today,
    }
    extra_global = [k for k, v in replacements.items() if v and k not in placeholders_in_doc]
    if extra_global:
        raise ValueError(
            "Плейсхолдеры отсутствуют в шаблоне: " + ", ".join(sorted(extra_global))
        )
    missing_global = [k for k in ("наименование закупки",) if k in placeholders_in_doc and not replacements[k]]
    if missing_global:
        raise ValueError(
            "Не заполнены обязательные поля: " + ", ".join(sorted(missing_global))
        )
    _replace_placeholders(doc, replacements)

    blocks = _find_position_blocks(doc)
    if not blocks:
        raise ValueError("В шаблоне не найдено ни одной секции с расчетом НМЦК и таблицей.")

    prototype_paragraph, prototype_table = blocks[0]
    paragraph_template_el = deepcopy(prototype_paragraph._element)
    table_template_el = deepcopy(prototype_table._element)
    message_paragraph_template_el = deepcopy(_get_message_paragraph_template(doc)._element)

    for extra_paragraph, extra_table in blocks[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)
        extra_table._element.getparent().remove(extra_table._element)

    pos_keys = set(_collect_placeholders_from_paragraph(prototype_paragraph))
    needs_heading = "предмет закупки" in pos_keys

    first_position = payload.positions[0]
    if not first_position.items:
        message_text = _build_no_items_message(first_position)
        message_el = deepcopy(message_paragraph_template_el)
        _insert_before(prototype_table._element, message_el)
        message_paragraph = Paragraph(message_el, doc)
        message_paragraph.text = message_text
        prototype_paragraph._element.getparent().remove(prototype_paragraph._element)
        prototype_table._element.getparent().remove(prototype_table._element)
        insert_after_element = message_el
    else:
        if needs_heading and not first_position.positionName:
            raise ValueError("positionName обязателен, так как используется в шаблоне.")
        if not needs_heading and first_position.positionName:
            raise ValueError("positionName передан, но отсутствует в шаблоне.")
        _set_position_heading_text(prototype_paragraph, first_position.positionName)
        _fill_items_table(
            prototype_table, first_position.items, first_position.positionPrice
        )
        insert_after_element = prototype_table._element

    for position in payload.positions[1:]:
        if not position.items:
            message_text = _build_no_items_message(position)
            message_el = deepcopy(message_paragraph_template_el)
            _insert_after(insert_after_element, message_el)
            message_paragraph = Paragraph(message_el, doc)
            message_paragraph.text = message_text
            insert_after_element = message_el
            continue

        new_paragraph_el = deepcopy(paragraph_template_el)
        _insert_after(insert_after_element, new_paragraph_el)
        new_paragraph = Paragraph(new_paragraph_el, doc)
        if needs_heading and not position.positionName:
            raise ValueError("positionName обязателен, так как используется в шаблоне.")
        if not needs_heading and position.positionName:
            raise ValueError("positionName передан, но отсутствует в шаблоне.")
        _set_position_heading_text(new_paragraph, position.positionName)

        new_table_el = deepcopy(table_template_el)
        _insert_after(new_paragraph_el, new_table_el)
        new_table = Table(new_table_el, doc)
        _fill_items_table(new_table, position.items, position.positionPrice)

        insert_after_element = new_table_el

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_docx_bytes(docx_bytes: bytes, target_ext: str) -> bytes:
    target_ext = target_ext.lower()
    if target_ext not in {"doc", "pdf"}:
        raise ValueError(f"Неподдерживаемый формат конвертации: {target_ext}")

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) не найден для конвертации документа.")

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        input_path = tmp_path / "document.docx"
        input_path.write_bytes(docx_bytes)

        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            target_ext,
            "--outdir",
            tmp_dir,
            str(input_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"Ошибка конвертации в {target_ext}: {result.stderr.strip() or result.stdout.strip()}"
            )

        output_candidates = list(tmp_path.glob(f"document.{target_ext}"))
        if not output_candidates:
            output_candidates = list(tmp_path.glob(f"*.{target_ext}"))
        if not output_candidates:
            raise RuntimeError(f"Конвертация в {target_ext} не создала файл.")

        return output_candidates[0].read_bytes()


app = FastAPI(title="TenderHack Price Justification")


@app.post("/api/v1/ste-price-justification/doc")
async def generate_ste_price_justification(payload: StePriceTemplateRequest) -> Response:
    if not payload.positions:
        raise HTTPException(status_code=400, detail="positions не должен быть пустым")
    for idx, position in enumerate(payload.positions, start=1):
        if not position.items:
            if not position.positionName or position.positionPrice is None:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"positions[{idx}].items пуст, поэтому positionName и positionPrice обязательны"
                    ),
                )

    try:
        docx_content = build_ste_price_docx_from_template(payload)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    doc_type = (payload.docType or "docx").lower()
    if doc_type not in {"docx", "doc", "pdf"}:
        raise HTTPException(status_code=400, detail="docType должен быть docx, doc или pdf")

    if doc_type == "docx":
        filename = "ste_price_justification.docx"
        headers = {"Content-Disposition": f"attachment; filename={filename}"}
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    try:
        converted = convert_docx_bytes(docx_content, doc_type)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    if doc_type == "doc":
        media_type = "application/msword"
        filename = "ste_price_justification.doc"
    else:
        media_type = "application/pdf"
        filename = "ste_price_justification.pdf"

    headers = {"Content-Disposition": f"attachment; filename={filename}"}
    return Response(content=converted, media_type=media_type, headers=headers)


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
